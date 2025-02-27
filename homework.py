import streamlit as st
from crewai import Agent, Task, Crew, Process
import os
from dotenv import load_dotenv
from langchain_openai import ChatOpenAI
from docx import Document
from io import BytesIO
import base64

load_dotenv()

os.environ["OPENAI_API_KEY"] = "your_openai_api_key" # Load environment variables for API Keys

def get_download_link(bio, filename):
    b64 = base64.b64encode(bio.read()).decode()
    return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">Download Homework Help Document</a>'

def generate_docx(content):
    doc = Document()
    doc.add_heading('Homework Help Document', 0)
    doc.add_paragraph(content)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

st.set_page_config(layout="wide")
st.title("AI Homework Helper")

# User Inputs for Homework Help
subject = st.selectbox('Select Subject', ('Mathematics', 'Science', 'History', 'Literature'))

question = st.text_area('Enter your homework question', 'Type here...')

additional_info = st.text_area('Any additional information', 'Type here...')

# Initialize Tools and LLM
# Here we can use any opensource LLM like claude as well

llm = ChatOpenAI(
    model="gpt-4o-mini",
    temperature=0.5
)

# Define Agents
homework_helper = Agent(
    role="Homework Helper",
    goal="Provide detailed answers and explanations to homework questions.",
    backstory="An AI trained across various subjects to assist students with their homework.",
    verbose=True,
    allow_delegation=True,
    llm=llm
)

# Define Task
homework_task = Task(
    description=(
        "1. Understand the student's question: {question}.\n"
        "2. Provide a detailed answer with explanations and, if applicable, step-by-step solutions.\n"
        "3. Include useful tips or study advice related to the question."
    ),
    expected_output="A comprehensive answer with explanations tailored to the student's needs.",
    agent=homework_helper
)

# Create Crew
crew = Crew(
    agents=[homework_helper],
    tasks=[homework_task],
    verbose=True
)

# Execution
if st.button("Get Help With Homework"):
    with st.spinner('Finding the best answers...'):
        results = crew.kickoff(inputs={"question": question, "additional_info": additional_info})
        st.write(results)
        docx_file = generate_docx(results)
        download_link = get_download_link(docx_file, "Homework_Help_Document.docx")
        st.markdown(download_link, unsafe_allow_html=True)
